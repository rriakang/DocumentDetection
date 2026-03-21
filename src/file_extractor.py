"""PDF, DOCX, HWP/HWPX 파일에서 텍스트를 행 단위로 추출."""

import os
import re
import subprocess
import tempfile
import zipfile
import xml.etree.ElementTree as ET
from pathlib import Path
from typing import List, Optional


def extract_rows_from_file(file_path: str) -> List[str]:
    """파일 형식에 따라 적절한 추출기를 호출하여 행 텍스트 목록을 반환."""
    ext = Path(file_path).suffix.lower()

    if ext == ".pdf":
        return _extract_from_pdf(file_path)
    elif ext == ".docx":
        return _extract_from_docx(file_path)
    elif ext in (".hwpx", ".hwtx"):
        return _extract_from_hwpx(file_path)
    elif ext == ".hwp":
        return _extract_from_hwp(file_path)
    else:
        raise ValueError(f"지원하지 않는 파일 형식: {ext}")


def split_documents(rows: List[str]) -> List[dict]:
    """텍스트 행 목록에서 여러 문서가 합쳐져 있는 경우 문서별로 분리.

    Returns:
        [{"title": "문서제목", "rows": [행들]}, ...]
        문서가 1개뿐이면 단일 항목 리스트 반환.
    """
    # 문서 제목 키워드
    _DOC_KEYWORDS = ["신청서", "확인서", "동의서", "요청서", "계약서", "위임장", "서약서", "증명서", "보고서"]

    # 문서 경계 찾기
    boundaries = []  # [(index, title), ...]
    for i, row in enumerate(rows):
        text = row.strip()
        # 짧은 제목성 행에서만 매칭 (긴 본문에 포함된 경우 제외)
        if len(text) > 40:
            continue
        # 공백 제거 후 키워드 매칭 (예: "확 인 서" → "확인서")
        no_space = text.replace(" ", "")
        if any(no_space.endswith(kw) for kw in _DOC_KEYWORDS):
            boundaries.append((i, text))

    # 경계가 0~1개면 분리 불필요
    if len(boundaries) <= 1:
        title = boundaries[0][1] if boundaries else ""
        return [{"title": title, "rows": rows}]

    # 여러 문서로 분리
    documents = []
    for idx, (start, title) in enumerate(boundaries):
        if idx + 1 < len(boundaries):
            end = boundaries[idx + 1][0]
        else:
            end = len(rows)
        doc_rows = rows[start:end]
        documents.append({"title": title, "rows": doc_rows})

    return documents


def _extract_from_pdf(file_path: str) -> List[str]:
    """pdfplumber로 PDF에서 텍스트+테이블 추출."""
    import pdfplumber

    rows = []
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            # 테이블이 있으면 테이블 우선 추출
            tables = page.extract_tables()
            if tables:
                for table in tables:
                    for row in table:
                        if row:
                            cells = [c.strip() if c else "" for c in row]
                            text = "  ".join(c for c in cells if c)
                            if text.strip():
                                rows.append(text.strip())
            else:
                # 테이블 없으면 일반 텍스트 추출
                text = page.extract_text()
                if text:
                    for line in text.split("\n"):
                        line = line.strip()
                        if line:
                            rows.append(line)
    return rows


def _extract_from_docx(file_path: str) -> List[str]:
    """python-docx로 Word 문서에서 텍스트+테이블 추출."""
    from docx import Document

    doc = Document(file_path)
    rows = []

    for element in doc.element.body:
        tag = element.tag.split("}")[-1] if "}" in element.tag else element.tag

        if tag == "tbl":
            # 테이블 처리
            for table in doc.tables:
                if table._element is element:
                    for row in table.rows:
                        cells = [cell.text.strip() for cell in row.cells]
                        text = "  ".join(c for c in cells if c)
                        if text.strip():
                            rows.append(text.strip())
                    break
        elif tag == "p":
            # 일반 문단 처리
            for para in doc.paragraphs:
                if para._element is element:
                    text = para.text.strip()
                    if text:
                        rows.append(text)
                    break

    return rows


def _extract_from_hwpx(file_path: str) -> List[str]:
    """HWPX (ZIP + XML) 파일에서 텍스트 추출."""
    rows = []

    try:
        with zipfile.ZipFile(file_path, "r") as zf:
            # HWPX 내부의 section XML 파일들을 찾음
            section_files = sorted([
                n for n in zf.namelist()
                if n.startswith("Contents/") and n.endswith(".xml")
                and "section" in n.lower()
            ])

            if not section_files:
                # section 파일이 없으면 모든 XML에서 텍스트 추출
                section_files = [
                    n for n in zf.namelist()
                    if n.startswith("Contents/") and n.endswith(".xml")
                ]

            for sf in section_files:
                xml_data = zf.read(sf)
                rows.extend(_parse_hwpx_xml(xml_data))

    except zipfile.BadZipFile:
        raise ValueError("HWPX 파일이 손상되었습니다.")

    return rows


def _parse_hwpx_xml(xml_data: bytes) -> List[str]:
    """HWPX의 section XML에서 텍스트를 추출."""
    rows = []

    try:
        root = ET.fromstring(xml_data)
    except ET.ParseError:
        return rows

    # 모든 텍스트 노드를 수집
    # HWPX namespace 처리
    ns_pattern = re.compile(r"\{[^}]+\}")

    current_para = []
    for elem in root.iter():
        local_tag = ns_pattern.sub("", elem.tag)

        if local_tag == "t" and elem.text:
            current_para.append(elem.text)
        elif local_tag in ("p", "para") and current_para:
            text = "".join(current_para).strip()
            if text:
                rows.append(text)
            current_para = []

    # 마지막 문단 처리
    if current_para:
        text = "".join(current_para).strip()
        if text:
            rows.append(text)

    return rows


def _extract_from_hwp(file_path: str) -> List[str]:
    """구버전 HWP 파일에서 텍스트 추출.

    1차: olefile로 PrvText 스트림에서 미리보기 텍스트 추출
    2차: LibreOffice 변환 (설치되어 있을 경우)
    """
    # 1차 시도: olefile로 PrvText 추출
    rows = _extract_hwp_prvtext(file_path)
    if rows:
        return rows

    # 2차 시도: LibreOffice로 PDF 변환 후 추출
    lo_path = _find_libreoffice()
    if not lo_path:
        raise RuntimeError(
            "HWP 파일에서 텍스트를 추출할 수 없습니다.\n"
            "olefile로 PrvText 추출에 실패했으며 LibreOffice도 설치되어 있지 않습니다.\n"
            "설치: brew install --cask libreoffice (macOS)\n"
            "또는: sudo apt install libreoffice (Linux)"
        )

    with tempfile.TemporaryDirectory() as tmpdir:
        result = subprocess.run(
            [lo_path, "--headless", "--convert-to", "pdf",
             "--outdir", tmpdir, file_path],
            capture_output=True, timeout=60,
        )
        if result.returncode != 0:
            raise RuntimeError(f"HWP→PDF 변환 실패: {result.stderr.decode()}")
        pdf_files = list(Path(tmpdir).glob("*.pdf"))
        if not pdf_files:
            raise RuntimeError("HWP→PDF 변환 결과물을 찾을 수 없습니다.")
        return _extract_from_pdf(str(pdf_files[0]))


def _extract_hwp_prvtext(file_path: str) -> List[str]:
    """olefile을 사용하여 HWP의 PrvText 스트림에서 텍스트를 추출.

    셀 구조를 분석하여 <라벨><빈셀> 쌍을 개별 입력 필드로 분리한다.
    예: <과제개요><지원 기관><><사업명><> → "과제개요 (필수입력)", "지원 기관", "사업명"
    """
    try:
        import olefile
    except ImportError:
        return []

    try:
        if not olefile.isOleFile(file_path):
            return []

        ole = olefile.OleFileIO(file_path)
        try:
            if not ole.exists("PrvText"):
                return []

            prv_data = ole.openstream("PrvText").read()
            text = prv_data.decode("utf-16-le", errors="ignore")

            rows = []
            for line in text.split("\n"):
                line = line.strip()
                if not line:
                    continue

                # 원본 셀 목록 (빈 셀 포함)
                raw_cells = re.findall(r"<([^>]*)>", line)
                if not raw_cells:
                    continue

                # <라벨><빈셀> 쌍 패턴 감지: 3개 이상 셀이 있고 라벨+빈셀 쌍이 있는 경우
                if len(raw_cells) >= 4:
                    sub_fields = _split_label_empty_pairs(raw_cells)
                    if sub_fields:
                        rows.extend(sub_fields)
                        continue

                # 일반 행: 빈 셀 제거 후 합치기
                non_empty = [c.strip() for c in raw_cells if c.strip()]
                if non_empty:
                    rows.append("  ".join(non_empty))

            return rows
        finally:
            ole.close()
    except Exception:
        return []


def _split_label_empty_pairs(cells: List[str]) -> Optional[List[str]]:
    """셀 목록에서 <라벨><빈셀> 쌍을 감지하여 개별 행으로 분리.

    패턴: [섹션헤더, 라벨1, "", 라벨2, "", ...] → 섹션헤더, 라벨1, 라벨2, ...
    패턴: [부모라벨, 하위1, "", 하위2, ""] → 부모라벨 > 하위1, 부모라벨 > 하위2

    빈 셀이 라벨 뒤에 반복적으로 나오면 각 라벨이 입력 필드.
    """
    # 라벨+빈셀 쌍 개수 세기
    pair_count = 0
    i = 0
    while i < len(cells) - 1:
        label = cells[i].strip()
        next_cell = cells[i + 1].strip()
        if label and not next_cell:
            pair_count += 1
            i += 2
        else:
            i += 1

    # 쌍이 2개 이상이면 분리 대상
    if pair_count < 2:
        return None

    result = []
    # 첫 셀이 섹션 헤더인지 판단 (뒤에 빈셀이 안 오는 경우)
    header = None
    start = 0
    first = cells[0].strip()
    if first and (len(cells) < 2 or cells[1].strip()):
        # 첫 셀 다음이 빈셀이 아님 → 섹션 헤더
        header = first
        result.append(header)
        start = 1
    elif first and not cells[1].strip():
        # 첫 셀+빈셀 쌍 → 첫 셀도 필드
        pass

    # 라벨+빈셀 쌍을 개별 행으로 추출
    i = start
    while i < len(cells):
        label = cells[i].strip()
        if not label:
            i += 1
            continue

        # 다음 셀이 빈셀이면 이 라벨은 입력 필드
        if i + 1 < len(cells) and not cells[i + 1].strip():
            result.append(label)
            i += 2
        else:
            # 빈셀이 안 따라오면 일반 텍스트
            result.append(label)
            i += 1

    return result if len(result) > 1 else None


def _find_libreoffice() -> Optional[str]:
    """시스템에 설치된 LibreOffice 경로를 찾는다."""
    candidates = [
        "/Applications/LibreOffice.app/Contents/MacOS/soffice",  # macOS
        "/usr/bin/libreoffice",  # Linux
        "/usr/bin/soffice",  # Linux alternative
        "/snap/bin/libreoffice",  # Snap
    ]

    for path in candidates:
        if os.path.isfile(path):
            return path

    # PATH에서 찾기
    try:
        result = subprocess.run(
            ["which", "libreoffice"],
            capture_output=True, text=True, timeout=5,
        )
        if result.returncode == 0 and result.stdout.strip():
            return result.stdout.strip()
    except Exception:
        pass

    return None
