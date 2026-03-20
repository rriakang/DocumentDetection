"""파일에 사용자 입력값을 채워넣는 모듈.

지원 형식:
- DOCX: 테이블 셀에서 필드명을 찾아 옆 셀에 값 삽입
- HWPX: XML에서 필드명 뒤에 값 삽입
- PDF: 지원 안 함 (읽기 전용)
"""

import os
import re
import shutil
import tempfile
import zipfile
import xml.etree.ElementTree as ET
from pathlib import Path
from typing import Dict, List, Optional


def fill_document(
    file_path: str,
    field_infos: List[dict],
    user_inputs: Dict[int, str],
) -> Optional[str]:
    """원본 파일을 복사하고 사용자 입력값을 채워서 새 파일 경로 반환.

    Args:
        file_path: 원본 파일 경로
        field_infos: 필드 정보 리스트 (id, field_name_ko 등)
        user_inputs: {field_id: user_value} 매핑

    Returns:
        채워진 파일의 경로 (임시 디렉터리). 실패 시 None.
    """
    ext = Path(file_path).suffix.lower()

    # 필드명 → 값 매핑 생성
    field_map = _build_field_map(field_infos, user_inputs)
    if not field_map:
        return None

    if ext == ".docx":
        return _fill_docx(file_path, field_map)
    elif ext == ".hwpx":
        return _fill_hwpx(file_path, field_map)
    elif ext == ".pdf":
        # PDF는 직접 수정이 어려움 → 대신 채워진 DOCX를 생성
        return _generate_filled_docx(file_path, field_infos, user_inputs)
    else:
        return None


def _build_field_map(
    field_infos: List[dict],
    user_inputs: Dict[int, str],
) -> Dict[str, str]:
    """필드 한글명 → 사용자 입력값 매핑."""
    field_map = {}
    for f in field_infos:
        fid = f.get("id")
        if fid in user_inputs:
            ko_name = f.get("field_name_ko", "").strip()
            if ko_name:
                # '*' 제거한 이름도 매핑 (필드명에 *가 붙어있을 수 있음)
                clean_name = ko_name.replace("*", "").strip()
                field_map[ko_name] = user_inputs[fid]
                if clean_name != ko_name:
                    field_map[clean_name] = user_inputs[fid]
    return field_map


# ── DOCX 채우기 ──────────────────────────────────────────────────────

def _fill_docx(file_path: str, field_map: Dict[str, str]) -> Optional[str]:
    """DOCX 파일의 테이블 셀에 값을 채워넣음."""
    from docx import Document
    from docx.shared import Pt

    try:
        doc = Document(file_path)
    except Exception:
        return None

    filled_count = 0

    for table in doc.tables:
        for row in table.rows:
            cells = row.cells
            for i, cell in enumerate(cells):
                cell_text = cell.text.strip()

                # 셀 텍스트에서 필드명 매칭
                matched_key = _match_field_name(cell_text, field_map)
                if matched_key is None:
                    continue

                value = field_map[matched_key]

                # ── 체크박스 필드: □를 ☑로 변경 ──
                if "□" in cell_text:
                    new_text = _fill_checkbox(cell_text, value)
                    _replace_cell_text(cell, new_text)
                    filled_count += 1
                    continue

                # ── 인접 빈 셀에 값 입력 ──
                # 같은 행에서 오른쪽 셀이 비어있으면 거기에 입력
                if i + 1 < len(cells):
                    next_cell = cells[i + 1]
                    next_text = next_cell.text.strip()
                    # 빈 셀이거나 '-'만 있으면 여기에 입력
                    if not next_text or next_text == "-" or next_text == "":
                        _set_cell_value(next_cell, value)
                        filled_count += 1
                        continue

                # 같은 셀에 필드명만 있고 비어있는 경우 (예: "성명*  ")
                # → 필드명 뒤에 값 추가
                if _is_label_only(cell_text):
                    _append_to_cell(cell, f"  {value}")
                    filled_count += 1

    # "신청인:" 라인에 이름 채우기
    name_value = None
    for key, val in field_map.items():
        if "성명" in key or "이름" in key:
            name_value = val
            break

    if name_value:
        for para in doc.paragraphs:
            if "신청인:" in para.text:
                old_text = para.text
                new_text = old_text.replace(
                    "신청인:                    ",
                    f"신청인:  {name_value}          ",
                )
                if new_text != old_text:
                    for run in para.runs:
                        if "신청인:" in run.text:
                            run.text = run.text.replace(
                                "신청인:                    ",
                                f"신청인:  {name_value}          ",
                            )
                            filled_count += 1
                            break

    if filled_count == 0:
        return None

    # 임시 파일로 저장
    stem = Path(file_path).stem
    tmp_dir = tempfile.mkdtemp()
    output_path = os.path.join(tmp_dir, f"{stem}_filled.docx")
    doc.save(output_path)
    return output_path


def _match_field_name(cell_text: str, field_map: Dict[str, str]) -> Optional[str]:
    """셀 텍스트에서 필드명을 매칭. 매칭된 키 반환."""
    cell_clean = cell_text.strip()
    if not cell_clean:
        return None

    # 정확 매칭
    for key in field_map:
        if key == cell_clean or key + "*" == cell_clean or cell_clean.startswith(key):
            return key

    # 부분 매칭 (셀에 필드명이 포함된 경우)
    for key in field_map:
        clean_key = key.replace("*", "").strip()
        if len(clean_key) >= 2 and clean_key in cell_clean:
            return key

    return None


def _is_label_only(text: str) -> bool:
    """셀이 필드 라벨만 포함하는지 (값이 없는지) 판별."""
    # "성명*", "주소*" 같은 패턴
    clean = text.replace("*", "").strip()
    return len(clean) <= 6 and "□" not in text


def _fill_checkbox(text: str, value: str) -> str:
    """체크박스 텍스트에서 사용자 선택값에 해당하는 □를 ☑로 변경."""
    # "□ 서울·인천·경기  □ 대전·충남·충북·세종" 형태
    parts = re.split(r"(□)", text)
    result = []
    for i, part in enumerate(parts):
        if part == "□" and i + 1 < len(parts):
            option_text = parts[i + 1].strip().split("□")[0].strip()
            # 사용자 입력값이 옵션 텍스트에 포함되어 있으면 체크
            if value.strip() in option_text or option_text in value.strip():
                result.append("☑")
            else:
                result.append("□")
        else:
            result.append(part)
    return "".join(result)


def _replace_cell_text(cell, new_text: str):
    """셀의 전체 텍스트를 교체."""
    for para in cell.paragraphs:
        if para.runs:
            para.runs[0].text = new_text
            for run in para.runs[1:]:
                run.text = ""
            return
    # runs가 없으면 직접 설정
    if cell.paragraphs:
        cell.paragraphs[0].text = new_text


def _set_cell_value(cell, value: str):
    """빈 셀에 값을 설정."""
    if cell.paragraphs:
        para = cell.paragraphs[0]
        if para.runs:
            para.runs[0].text = value
        else:
            para.text = value
    else:
        cell.text = value


def _append_to_cell(cell, suffix: str):
    """셀의 마지막 run에 텍스트를 추가."""
    for para in cell.paragraphs:
        if para.runs:
            para.runs[-1].text += suffix
            return
    if cell.paragraphs:
        cell.paragraphs[0].text += suffix


# ── HWPX 채우기 ──────────────────────────────────────────────────────

def _fill_hwpx(file_path: str, field_map: Dict[str, str]) -> Optional[str]:
    """HWPX (ZIP+XML) 파일에서 필드명 뒤에 값을 삽입."""
    try:
        stem = Path(file_path).stem
        tmp_dir = tempfile.mkdtemp()
        output_path = os.path.join(tmp_dir, f"{stem}_filled.hwpx")

        # 원본 ZIP 복사
        shutil.copy2(file_path, output_path)

        filled_count = 0

        with zipfile.ZipFile(output_path, "r") as zf_read:
            namelist = zf_read.namelist()
            section_files = [
                n for n in namelist
                if n.startswith("Contents/") and n.endswith(".xml")
                and "section" in n.lower()
            ]
            if not section_files:
                section_files = [
                    n for n in namelist
                    if n.startswith("Contents/") and n.endswith(".xml")
                ]

            # 각 section XML을 수정
            modified_files = {}
            for sf in section_files:
                xml_data = zf_read.read(sf)
                new_xml, count = _fill_hwpx_xml(xml_data, field_map)
                if count > 0:
                    modified_files[sf] = new_xml
                    filled_count += count

        if filled_count == 0:
            # 아무것도 못 채웠으면 None
            shutil.rmtree(tmp_dir, ignore_errors=True)
            return None

        # 수정된 XML로 ZIP 재구성
        tmp_output = output_path + ".tmp"
        with zipfile.ZipFile(output_path, "r") as zf_read:
            with zipfile.ZipFile(tmp_output, "w", zipfile.ZIP_DEFLATED) as zf_write:
                for item in zf_read.namelist():
                    if item in modified_files:
                        zf_write.writestr(item, modified_files[item])
                    else:
                        zf_write.writestr(item, zf_read.read(item))

        os.replace(tmp_output, output_path)
        return output_path

    except Exception:
        return None


def _fill_hwpx_xml(xml_data: bytes, field_map: Dict[str, str]) -> tuple:
    """HWPX XML에서 필드명을 찾아 값을 삽입. (수정된 XML bytes, 채운 개수) 반환."""
    try:
        # XML 문자열로 처리 (namespace 보존)
        xml_str = xml_data.decode("utf-8")
        filled_count = 0

        for field_name, value in field_map.items():
            clean_name = field_name.replace("*", "").strip()
            if len(clean_name) < 2:
                continue

            # 체크박스 필드: □를 ☑로 변경
            if "□" in xml_str and value:
                # 필드명이 포함된 라인에서 체크박스 처리
                pattern = re.compile(
                    rf"(<hp:t>)([^<]*{re.escape(clean_name)}[^<]*□[^<]*)(</hp:t>)"
                )
                match = pattern.search(xml_str)
                if match:
                    original = match.group(2)
                    filled = _fill_checkbox(original, value)
                    if filled != original:
                        xml_str = xml_str.replace(
                            match.group(0),
                            f"{match.group(1)}{filled}{match.group(3)}",
                        )
                        filled_count += 1
                        continue

            # 일반 필드: "필드명*" 뒤에 값 추가
            # 패턴1: <hp:t>성명*</hp:t> → <hp:t>성명*  홍길동</hp:t>
            for pattern_name in [field_name, clean_name, clean_name + "*"]:
                escaped = re.escape(pattern_name)
                pattern = re.compile(
                    rf"(<hp:t>)({escaped}\s*)(</hp:t>)"
                )
                match = pattern.search(xml_str)
                if match:
                    original = match.group(2)
                    xml_str = xml_str.replace(
                        match.group(0),
                        f"{match.group(1)}{original}  {value}{match.group(3)}",
                    )
                    filled_count += 1
                    break

            # 패턴2: "신청인:" 처리
            if "신청인" in clean_name or "성명" in clean_name:
                applicant_pattern = re.compile(
                    r"(<hp:t>)(신청인:\s*)(</hp:t>)"
                )
                match = applicant_pattern.search(xml_str)
                if match:
                    xml_str = xml_str.replace(
                        match.group(0),
                        f"{match.group(1)}신청인:  {value}          {match.group(3)}",
                    )

        return xml_str.encode("utf-8"), filled_count

    except Exception:
        return xml_data, 0


# ── PDF용 대체: 채워진 DOCX 생성 ──────────────────────────────────────

def _generate_filled_docx(
    original_path: str,
    field_infos: List[dict],
    user_inputs: Dict[int, str],
) -> Optional[str]:
    """PDF 원본 대신 입력값이 채워진 새 DOCX를 생성."""
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    try:
        doc = Document()

        # 제목
        stem = Path(original_path).stem
        title = doc.add_heading(f"{stem} - 작성 완료", level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph("")

        # 필드 테이블 생성
        table = doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"

        # 헤더
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "항목"
        hdr_cells[1].text = "입력값"

        for f in field_infos:
            fid = f.get("id")
            if fid not in user_inputs:
                continue

            ko = f.get("field_name_ko", "")
            en = f.get("field_name", ko)
            display = f"{ko} ({en})" if ko and ko != en else (ko or en)

            row_cells = table.add_row().cells
            row_cells[0].text = display
            row_cells[1].text = user_inputs[fid]

        # 저장
        tmp_dir = tempfile.mkdtemp()
        output_path = os.path.join(tmp_dir, f"{stem}_filled.docx")
        doc.save(output_path)
        return output_path

    except Exception:
        return None
